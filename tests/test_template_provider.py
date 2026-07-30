"""The uploaded document as a knowledge source.

Parsing and fidelity are covered by `test_docx_parse.py` and `test_docx_fidelity.py`. What
is tested here is the seam: where the bytes go, how the provider finds them, what it
contributes, and — most importantly — what it refuses.
"""

from __future__ import annotations

import json
import uuid
from collections.abc import AsyncIterator
from pathlib import Path

import pytest
import pytest_asyncio
from sqlalchemy import delete
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from backend.artifacts import Artifact, ArtifactStore
from backend.core.config import settings
from backend.core.run_context import RunContext
from backend.invariants.docx_parse import TemplateError
from backend.knowledge.providers.template import POINTER_PATH, TemplateProvider, store_template
from backend.phase_a.gather import gather
from backend.phase_a.resolver import plan_for
from backend.schemas.errors import WorkspaceError
from backend.schemas.intent import IntentObject
from backend.schemas.template import TemplateObject
from backend.storage import get_storage
from backend.workspace.models import Contract
from backend.workspace.store import WorkspaceStore

FIXTURE = Path(__file__).parent / "data" / "sla-sample.docx"
INTENT = IntentObject(contract_type="sla", confidence=0.9, mode="template")


@pytest.fixture(scope="module")
def docx() -> bytes:
    return FIXTURE.read_bytes()


@pytest_asyncio.fixture
async def ctx() -> AsyncIterator[RunContext]:
    engine = create_async_engine(settings.database_url)
    factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
    cid = uuid.uuid4()
    async with factory() as s:
        s.add(Contract(id=cid, contract_type="sla", request="Convert this SLA"))
        await s.commit()
    try:
        yield RunContext(contract_id=cid, session_factory=factory, contract_type="sla")
    finally:
        async with factory() as s:
            await s.execute(delete(Contract).where(Contract.id == cid))
            await s.commit()
        await engine.dispose()


# ---------------------------------------------------------------------------- the seam


async def test_the_binary_goes_to_storage_and_only_a_pointer_to_the_workspace(
    ctx: RunContext, docx: bytes
) -> None:
    """`workspace_files.content` is a Text column. A 10 MB DOCX cannot live there, and
    would be loaded by every listing that touches the row."""
    key = await store_template(docx, "sla-sample.docx", ctx)

    assert get_storage().exists(key)
    assert get_storage().get(key) == docx

    async with ctx.session_factory() as s:
        pointer = json.loads(await WorkspaceStore(s).read(ctx.contract_id, POINTER_PATH))
    assert pointer == {"storage_key": key, "filename": "sla-sample.docx"}
    assert len(pointer["storage_key"]) < 60, "a pointer, not a payload"


async def test_the_pointer_lives_apart_from_reference_documents(
    ctx: RunContext, docx: bytes
) -> None:
    """Template text is authoritative; reference text must never be copied. The two paths
    are opposites, so they never share a prefix."""
    await store_template(docx, "sla-sample.docx", ctx)
    assert POINTER_PATH.startswith("template/")
    assert not POINTER_PATH.startswith("references/")


# ------------------------------------------------------------------------- availability


async def test_a_run_with_no_upload_does_not_get_the_template_provider(
    ctx: RunContext,
) -> None:
    assert await TemplateProvider().available(INTENT, ctx) is False

    plan = await plan_for(INTENT, ctx)
    assert "template" not in plan.providers


async def test_a_run_with_an_upload_does(ctx: RunContext, docx: bytes) -> None:
    await store_template(docx, "sla-sample.docx", ctx)

    assert await TemplateProvider().available(INTENT, ctx) is True

    plan = await plan_for(INTENT, ctx)
    assert "template" in plan.providers
    assert plan.providers.index("template") < plan.providers.index("llm")


# ------------------------------------------------------------------------ contributing


async def test_the_provider_contributes_formatting_and_writes_the_artifact(
    ctx: RunContext, docx: bytes
) -> None:
    await store_template(docx, "sla-sample.docx", ctx)

    contribution = await TemplateProvider().contribute(INTENT, ctx)

    assert contribution.provider == "template"
    assert contribution.formatting is not None
    assert len(contribution.formatting.blocks) > 10
    assert contribution.formatting.tables, "the fee table must survive into the manifest"

    stored = await ArtifactStore(ctx.session_factory, ctx.contract_id).load(Artifact.TEMPLATE)
    assert isinstance(stored, TemplateObject)
    assert stored.filename == "sla-sample.docx"


async def test_it_contributes_structure_but_claims_no_meaning(ctx: RunContext, docx: bytes) -> None:
    """Parsing has a right answer; interpretation does not.

    A parser's guess about what a heading means would otherwise outrank the understanding
    agent's reading of it, because `template` sits above nothing but `reference` and `llm`
    — but it sits above them.
    """
    await store_template(docx, "sla-sample.docx", ctx)

    contribution = await TemplateProvider().contribute(INTENT, ctx)

    assert contribution.sections == ()
    assert contribution.clause_candidates == ()


async def test_it_participates_in_a_real_gather(ctx: RunContext, docx: bytes) -> None:
    await store_template(docx, "sla-sample.docx", ctx)

    plan = await plan_for(INTENT, ctx)
    contributions = await gather(plan, INTENT, ctx)

    # The playbook is always available, so it participates too; assert the template's
    # position relative to the floor rather than the exact set, which grows as providers
    # ship. Precedence order, not completion order.
    providers = [c.provider for c in contributions]
    assert "template" in providers
    assert providers.index("template") < providers.index("llm")
    assert providers[-1] == "llm", "the floor stays last whatever else participates"


# --------------------------------------------------------------------------- refusals


async def test_an_unusable_upload_is_refused_at_intake_not_mid_run(ctx: RunContext) -> None:
    """Refuse while the user is still watching.

    Deferring the parse to contribute time means the user waits through intent
    determination and provider resolution before being told the file was never usable.
    """
    from docx import Document
    from docx.oxml.ns import qn

    document = Document()
    paragraph = document.add_paragraph("Term: ")
    ins = paragraph._element.makeelement(qn("w:ins"), {})
    run = ins.makeelement(qn("w:r"), {})
    node = run.makeelement(qn("w:t"), {})
    node.text = "thirty days"
    run.append(node)
    ins.append(run)
    paragraph._element.append(ins)

    from io import BytesIO

    buf = BytesIO()
    document.save(buf)

    with pytest.raises(TemplateError, match="tracked changes"):
        await store_template(buf.getvalue(), "tracked.docx", ctx)

    async with ctx.session_factory() as s:
        assert await WorkspaceStore(s).exists(ctx.contract_id, POINTER_PATH) is False


async def test_a_pointer_whose_blob_is_gone_is_refused(ctx: RunContext, docx: bytes) -> None:
    """Not silently drafting from nothing.

    Template mode without a formatting manifest regenerates the document, losing exactly
    what the upload was for — so a missing blob must stop the run, not degrade it.
    """
    await store_template(docx, "sla-sample.docx", ctx)

    async with ctx.session_factory() as s:
        await WorkspaceStore(s).write(
            ctx.contract_id,
            POINTER_PATH,
            json.dumps({"storage_key": f"{uuid.uuid4()}.docx", "filename": "gone.docx"}),
        )
        await s.commit()

    with pytest.raises(WorkspaceError, match="missing from storage"):
        await TemplateProvider().contribute(INTENT, ctx)


async def test_storage_keys_are_opaque_and_flat(ctx: RunContext, docx: bytes) -> None:
    """`LocalStorage` refuses anything that could escape its root, so a key derived from
    the user's filename would be both a traversal surface and a collision."""
    key = await store_template(docx, "../../etc/passwd.docx", ctx)

    assert "/" not in key and ".." not in key
    assert key.endswith(".docx")
