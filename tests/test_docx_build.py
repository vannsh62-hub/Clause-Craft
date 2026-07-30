"""The document renderer: what gets numbered, what does not, and byte stability.

These assertions are the whole difference between a generated draft and something that reads
as a contract, so each one names the defect it prevents rather than the property it checks.
"""

from __future__ import annotations

import io

import docx

from backend.invariants.docx_build import document_to_text, render_contract
from backend.schemas.document import ClauseNode, ContractDocument, Signatory

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

DOC = ContractDocument(
    title="Services Agreement",
    preamble='This Services Agreement is made on 2026-08-01 between Nova (the "Provider").',
    recitals=("WHEREAS the Client wishes to procure the Services;",),
    clauses=(
        ClauseNode(
            heading="Definitions",
            text="In this Agreement:",
            children=(ClauseNode(heading="Services", text="means the services in Schedule 1."),),
        ),
        ClauseNode(heading="Fees and Payment", text="Payment falls due within 30 days."),
    ),
    execution_note="IN WITNESS WHEREOF the parties have executed this Agreement.",
    signatories=(Signatory(name="Nova Techset Pvt Ltd", role="Service Provider"),),
)


def _paragraphs(payload: bytes) -> list[docx.text.paragraph.Paragraph]:
    return list(docx.Document(io.BytesIO(payload)).paragraphs)


def _is_numbered(paragraph: docx.text.paragraph.Paragraph) -> bool:
    return paragraph._p.find(f".//{_W}numPr") is not None


def test_clauses_are_numbered_by_word_not_typed_into_the_text() -> None:
    """A typed "1.1" becomes a second, wrong number the moment a clause is inserted."""
    payload = render_contract(DOC)
    numbered = [p.text for p in _paragraphs(payload) if _is_numbered(p)]

    assert numbered == ["Definitions", "Services", "Fees and Payment"]
    for text in numbered:
        assert not text[0].isdigit(), f"{text!r} carries a hand-typed number"


def test_the_preamble_recitals_and_signature_block_are_not_numbered() -> None:
    """A numbered WHEREAS is the clearest tell that a machine assembled the document."""
    payload = render_contract(DOC)
    by_text = {p.text.strip(): p for p in _paragraphs(payload) if p.text.strip()}

    for line in (
        'This Services Agreement is made on 2026-08-01 between Nova (the "Provider").',
        "WHEREAS the Client wishes to procure the Services;",
        "IN WITNESS WHEREOF the parties have executed this Agreement.",
        "Name:",
    ):
        assert line in by_text, f"{line!r} is missing from the document"
        assert not _is_numbered(by_text[line]), f"{line!r} must not be numbered"


def test_the_signature_block_starts_on_a_fresh_page() -> None:
    payload = render_contract(DOC)
    xml = docx.Document(io.BytesIO(payload)).element.xml
    assert 'w:type="page"' in xml, "the execution page must follow a page break"


def test_the_numbering_definition_is_installed() -> None:
    """A `w:numPr` pointing at an id with no `w:num` renders as no numbering at all."""
    payload = render_contract(DOC)
    numbering = docx.Document(io.BytesIO(payload)).part.numbering_part.element.xml
    assert 'w:abstractNumId w:val="9100"' in numbering or "9100" in numbering


def test_the_same_document_renders_to_the_same_bytes() -> None:
    """`exports.sha256` is only meaningful while this holds."""
    assert render_contract(DOC) == render_contract(DOC)


def test_the_text_form_carries_everything_the_docx_carries() -> None:
    """The gates read the text form; what they cannot see, they cannot refuse."""
    text = document_to_text(DOC)

    assert "# Services Agreement" in text
    assert "is made on 2026-08-01" in text
    assert "WHEREAS" in text
    assert "1. Definitions" in text
    assert "1.1 Services" in text
    assert "2. Fees and Payment" in text
    assert "IN WITNESS WHEREOF" in text
    assert "Nova Techset Pvt Ltd" in text


def test_an_empty_document_still_renders() -> None:
    """A title and nothing else must not raise — a blocked run still needs its draft kept."""
    payload = render_contract(ContractDocument(title="Agreement"))
    assert payload[:2] == b"PK"
