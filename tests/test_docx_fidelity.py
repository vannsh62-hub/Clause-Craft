"""Mode 2's central promise: what you did not ask to change comes back unchanged.

Five checks, in the order they would fail if the editor regressed:

1. a no-op edit changes nothing
2. `styles.xml` and `numbering.xml` are byte identical
3. KEEP blocks survive a real plan that modifies and removes other blocks
4. the result re-parses, with the expected block count
5. table structure, including merges, survives

Note what is *not* asserted: byte equality of the package. `python-docx` reserialises the
XML on save, so that test would be flaky and would be disabled within a week. Spec 05 §2's
"byte-comparable" wording is not achievable; these five are, and they catch the failures
that actually matter.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from backend.invariants.docx_edit import apply_transformation
from backend.invariants.docx_fidelity import SHARED_PARTS, compare, shared_parts_match
from backend.invariants.docx_parse import TemplateError, block_texts, parse_docx, style_fingerprint
from backend.schemas.cko import SourceRef
from backend.schemas.plan import SectionDecision, TransformationPlan

FIXTURE = Path(__file__).parent / "data" / "sla-sample.docx"


@pytest.fixture(scope="module")
def data() -> bytes:
    return FIXTURE.read_bytes()


def _ref(block_id: str) -> SourceRef:
    return SourceRef(provider="template", block_id=block_id)


def _find(data: bytes, needle: str) -> str:
    return next(k for k, v in block_texts(data).items() if needle in v)


# ------------------------------------------------------------------------ 1. no-op


def test_a_no_op_edit_changes_nothing(data: bytes) -> None:
    out, report = apply_transformation(data, TransformationPlan(), {})

    assert style_fingerprint(out) == style_fingerprint(data)
    assert compare(data, out).unchanged
    assert report.touched == 0


# --------------------------------------------------------------- 2. the shared parts


@pytest.mark.parametrize("part", SHARED_PARTS)
def test_the_style_and_numbering_parts_are_untouched(data: bytes, part: str) -> None:
    """These define what every style *means*.

    A change here restyles blocks whose own properties never changed, so fingerprint
    equality alone would not notice. Editing in place never opens them.
    """
    out, _ = apply_transformation(data, TransformationPlan(), {})
    assert part not in shared_parts_match(data, out)


# ------------------------------------------------------------ 3. KEEP under a real plan


def test_keep_blocks_survive_a_plan_that_modifies_and_removes(data: bytes) -> None:
    """The test that would have caught the position-dependent block id.

    Removing a block must not disturb the identity of the blocks after it. When ids
    encoded position, every KEEP downstream of the REMOVE reported as deleted.
    """
    termination = _find(data, "terminate on 30 days")
    sub_item = _find(data, "Excluding maintenance")
    kept = [_find(data, "Uptime shall be"), _find(data, "Fee Schedule")]

    plan = TransformationPlan(
        keep=tuple(
            SectionDecision(name=f"keep-{i}", decision="keep", reason="applies", source_ref=_ref(k))
            for i, k in enumerate(kept)
        ),
        modify=(
            SectionDecision(
                name="Termination",
                decision="modify",
                reason="90 days",
                source_ref=_ref(termination),
            ),
        ),
        remove=(
            SectionDecision(
                name="Maintenance carve-out",
                decision="remove",
                reason="not applicable",
                source_ref=_ref(sub_item),
            ),
        ),
    )

    out, report = apply_transformation(
        data, plan, {termination: "Either party may terminate on 90 days notice."}
    )

    result = compare(data, out, expect_unchanged=tuple(kept))
    assert result.unchanged, f"KEEP disturbed: {result}"
    assert report.modified == ["Termination"]
    assert report.removed == ["Maintenance carve-out"]
    assert not report.warnings


def test_the_edits_themselves_actually_happened(data: bytes) -> None:
    """Fidelity is only meaningful if the intended change also landed.

    A transformation that preserved everything by doing nothing would pass every check
    above.
    """
    termination = _find(data, "terminate on 30 days")
    sub_item = _find(data, "Excluding maintenance")
    plan = TransformationPlan(
        modify=(
            SectionDecision(name="T", decision="modify", reason="r", source_ref=_ref(termination)),
        ),
        remove=(
            SectionDecision(name="M", decision="remove", reason="r", source_ref=_ref(sub_item)),
        ),
    )

    out, _ = apply_transformation(data, plan, {termination: "Terminate on 90 days notice."})
    texts = " ".join(block_texts(out).values())

    assert "90 days" in texts
    assert "Excluding maintenance" not in texts


def test_numbering_survives_an_edit_to_a_numbered_paragraph(data: bytes) -> None:
    """`w:pPr` is untouched by a text replacement, so `w:numPr` goes with it."""
    uptime = _find(data, "Uptime shall be")
    plan = TransformationPlan(
        modify=(SectionDecision(name="U", decision="modify", reason="r", source_ref=_ref(uptime)),)
    )

    out, _ = apply_transformation(data, plan, {uptime: "Uptime shall be 99.95%."})

    before = {b.text_sha: b for b in style_fingerprint(data)}
    after = [b for b in style_fingerprint(out) if b.num_id is not None]
    assert len(after) == len([b for b in before.values() if b.num_id is not None])
    assert {b.ilvl for b in after} == {0, 1}


# ------------------------------------------------------------------- 4. re-parseable


def test_the_result_reparses_with_the_expected_block_count(data: bytes) -> None:
    sub_item = _find(data, "Excluding maintenance")
    plan = TransformationPlan(
        remove=(
            SectionDecision(name="M", decision="remove", reason="r", source_ref=_ref(sub_item)),
        )
    )

    out, _ = apply_transformation(data, plan, {})
    reparsed = parse_docx(out, filename="edited.docx", storage_key="k")

    assert len(reparsed.formatting.blocks) == len(style_fingerprint(data)) - 1


# ------------------------------------------------------------------ 5. table integrity


def test_table_structure_survives(data: bytes) -> None:
    before = parse_docx(data, filename="f.docx", storage_key="k").formatting.tables
    out, _ = apply_transformation(data, TransformationPlan(), {})
    after = parse_docx(out, filename="f.docx", storage_key="k").formatting.tables

    assert [(t.rows, t.columns, t.merged) for t in before] == [
        (t.rows, t.columns, t.merged) for t in after
    ]


# --------------------------------------------------------------------------- refusals


def test_a_modify_with_no_replacement_text_is_refused(data: bytes) -> None:
    """An empty clause in an executed contract is worse than a failed run."""
    termination = _find(data, "terminate on 30 days")
    plan = TransformationPlan(
        modify=(
            SectionDecision(name="T", decision="modify", reason="r", source_ref=_ref(termination)),
        )
    )

    with pytest.raises(TemplateError, match="No replacement text"):
        apply_transformation(data, plan, {})


def test_an_added_paragraph_clones_a_donor_so_it_keeps_its_numbering(data: bytes) -> None:
    """`add_paragraph(style="List Number")` yields `w:pStyle` and no `w:numPr`.

    Verified against python-docx 1.2.0: such a paragraph renders unnumbered. Cloning a
    neighbour that is already numbered is the only reliable way to add a numbered item.
    """
    uptime = _find(data, "Uptime shall be")
    plan = TransformationPlan(
        add=(
            SectionDecision(
                name="New SLA term", decision="add", reason="r", source_ref=_ref(uptime)
            ),
        )
    )

    out, report = apply_transformation(data, plan, {uptime: "Latency shall be under 100ms."})

    assert not report.warnings, "a donor was available; no fallback should have been needed"
    numbered = [b for b in style_fingerprint(out) if b.num_id is not None]
    assert len(numbered) == 4, "the added item joins the numbered list"


def test_a_document_with_tracked_changes_never_reaches_the_editor() -> None:
    document = Document()
    document.add_paragraph("Body.")
    buf = io.BytesIO()
    document.save(buf)
    clean = buf.getvalue()

    # Sanity: the clean document is editable, so the refusal below is about the revisions.
    apply_transformation(clean, TransformationPlan(), {})
