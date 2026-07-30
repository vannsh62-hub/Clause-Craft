"""Parsing a real DOCX.

The fixture (`tests/data/sla-sample.docx`) deliberately contains the four structures that
break naive parsers: nested numbering with explicit `w:numPr`, a table with a merged
header row, a header and footer, and an annexure behind a section break. A parser that only
walks `document.paragraphs` gets several of these wrong and reports success.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from backend.invariants.docx_parse import (
    TemplateError,
    block_ids,
    block_texts,
    parse_docx,
    style_fingerprint,
)

FIXTURE = Path(__file__).parent / "data" / "sla-sample.docx"


@pytest.fixture(scope="module")
def data() -> bytes:
    return FIXTURE.read_bytes()


def _parse(data: bytes):  # type: ignore[no-untyped-def]
    return parse_docx(data, filename="sla-sample.docx", storage_key="templates/x.docx")


# ------------------------------------------------------------------------- structure


def test_the_document_parses_into_blocks(data: bytes) -> None:
    template = _parse(data)
    assert len(template.formatting.blocks) > 10
    assert template.sha256 and template.size_bytes == len(data)


def test_paragraphs_inside_tables_are_not_lost(data: bytes) -> None:
    """`document.paragraphs` omits them, which is where fee tables and schedules live.

    The fixture's fee table holds a `{{ rate }}` placeholder. A parser that missed table
    content would report the document as having no rate to fill in — and would then
    produce a contract with a blank fee.
    """
    template = _parse(data)
    tokens = {p.token for p in template.placeholders}
    assert "{{ rate }}" in tokens

    document = Document(io.BytesIO(data))
    assert not any("{{ rate }}" in p.text for p in document.paragraphs), (
        "if this fails the fixture changed and the test no longer proves anything"
    )


def test_nested_numbering_is_captured(data: bytes) -> None:
    template = _parse(data)
    numbered = [b for b in template.formatting.blocks if b.num_id is not None]

    assert len(numbered) == 3
    assert {b.ilvl for b in numbered} == {0, 1}, "both list levels must survive"
    assert template.formatting.numbering.max_depth == 2


def test_table_shape_including_merges_is_captured(data: bytes) -> None:
    template = _parse(data)
    assert len(template.formatting.tables) == 1
    table = template.formatting.tables[0]
    assert (table.rows, table.columns) == (3, 3)
    assert table.merged, "the merged header row must be recorded"


def test_headers_footers_and_section_breaks_are_noticed(data: bytes) -> None:
    formatting = _parse(data).formatting
    assert formatting.has_headers and formatting.has_footers
    assert formatting.section_breaks == 1, "the annexure break, not the body's closing sectPr"


@pytest.mark.parametrize("token", ["[PROVIDER]", "[CUSTOMER]", "{{ rate }}", "<INSERT SPEC>"])
def test_every_placeholder_convention_is_recognised(data: bytes, token: str) -> None:
    assert token in {p.token for p in _parse(data).placeholders}


# ------------------------------------------------------------------------- identity


def test_block_ids_are_stable_across_a_load_and_save(data: bytes) -> None:
    """A round trip through python-docx must not change any identity.

    Ids are what `TransformationPlan.source_ref` points at. If they moved on save, every
    KEEP decision would dangle after the first edit.
    """
    document = Document(io.BytesIO(data))
    out = io.BytesIO()
    document.save(out)

    assert block_ids(style_fingerprint(data)) == block_ids(style_fingerprint(out.getvalue()))


def test_block_ids_do_not_depend_on_position() -> None:
    """The property that makes ids survive an edit.

    An id containing the block index would renumber everything after a deletion, so the
    surviving blocks would look deleted and every later KEEP would point at nothing.
    """
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    document.save(buf)
    before = buf.getvalue()

    trimmed = Document(io.BytesIO(before))
    body = trimmed.element.body
    first = body.findall(qn("w:p"))[0]
    first.getparent().remove(first)
    buf2 = io.BytesIO()
    trimmed.save(buf2)

    ids_before = dict(
        zip(block_ids(style_fingerprint(before)), block_texts(before).values(), strict=True)
    )
    ids_after = block_ids(style_fingerprint(buf2.getvalue()))

    surviving = [i for i, text in ids_before.items() if "Second" in text]
    assert surviving and surviving[0] in ids_after


def test_identical_paragraphs_get_distinct_ids() -> None:
    """Contracts really do repeat paragraphs — blank ones, "Intentionally left blank"."""
    document = Document()
    for _ in range(3):
        document.add_paragraph("Intentionally left blank.")
    buf = io.BytesIO()
    document.save(buf)

    ids = block_ids(style_fingerprint(buf.getvalue()))
    assert len(ids) == len(set(ids))


# ------------------------------------------------------------------------- refusals


def test_a_document_with_tracked_changes_is_refused() -> None:
    """Neither accepting nor rejecting silently is defensible.

    Accepting mutates a legal document with no audit trail; rejecting discards negotiated
    edits. Either way the document that was reviewed is not the document that was parsed,
    so the fidelity guarantee cannot be made. The user resolves them in Word in seconds.
    """
    document = Document()
    paragraph = document.add_paragraph("Agreed term: ")
    ins = paragraph._element.makeelement(qn("w:ins"), {})
    run = ins.makeelement(qn("w:r"), {})
    node = run.makeelement(qn("w:t"), {})
    node.text = "thirty days"
    run.append(node)
    ins.append(run)
    paragraph._element.append(ins)

    buf = io.BytesIO()
    document.save(buf)

    with pytest.raises(TemplateError, match="tracked changes"):
        parse_docx(buf.getvalue(), filename="tracked.docx", storage_key="k")
