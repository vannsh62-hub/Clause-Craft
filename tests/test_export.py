"""DOCX export: byte stability, the watermark, and the gate.

The gate is the point. If a document could be produced from anything other than a finalized
version, every deterministic check upstream would be theatre.
"""

from __future__ import annotations

import io
import time
import uuid
import zipfile
from collections.abc import AsyncIterator
from pathlib import Path

import pytest
import pytest_asyncio
from docx import Document
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from backend.core.config import settings
from backend.core.run_context import RunContext
from backend.invariants.export import DISCLAIMER, WATERMARK, docx_sha256, markdown_to_docx
from backend.storage import LocalStorage
from backend.tools.export_tool import export_docx
from backend.tools.finalize_tool import finalize_contract
from backend.workspace.models import Contract, ContractVersion, Export
from tests.helpers import faithful_markdown, render_clauses_for, tool_ctx

SAMPLE = """# Non-Disclosure Agreement

This Agreement is made on 1 August 2026 between ABC Pvt Ltd and XYZ Pvt Ltd.

## Confidentiality

XYZ Pvt Ltd shall hold all Confidential Information in strict confidence.

## Execution

Name: Jane Rao
Signature: ______________________
Date: ______________________
"""


# ------------------------------------------------------------------ the bytes themselves


def test_the_same_markdown_always_produces_the_same_bytes() -> None:
    """python-docx stamps every zip entry with the wall clock, so two saves of identical
    content differ. `_normalise_zip` fixes the container; this proves it."""
    first = markdown_to_docx(SAMPLE, title="NDA")
    time.sleep(1.1)  # long enough for a DOS timestamp to tick
    second = markdown_to_docx(SAMPLE, title="NDA")

    assert docx_sha256(first) == docx_sha256(second)


def test_different_content_produces_different_bytes() -> None:
    a = markdown_to_docx(SAMPLE, title="NDA")
    b = markdown_to_docx(SAMPLE + "\n\nAnd one more clause.", title="NDA")
    assert docx_sha256(a) != docx_sha256(b)


def test_the_archive_is_a_valid_opc_package() -> None:
    payload = markdown_to_docx(SAMPLE, title="NDA")
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        assert archive.testzip() is None
        assert archive.namelist()[0] == "[Content_Types].xml", "OPC expects it first"
        assert "word/document.xml" in archive.namelist()
        assert {info.date_time for info in archive.infolist()} == {(1980, 1, 1, 0, 0, 0)}


def test_the_document_reopens_and_carries_the_contract(tmp_path: Path) -> None:
    path = tmp_path / "contract.docx"
    path.write_bytes(markdown_to_docx(SAMPLE, title="Non-Disclosure Agreement"))

    document = Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs)

    assert "Non-Disclosure Agreement" in text
    assert "strict confidence" in text
    assert "Signature: ______________________" in text, "the signature block must not reflow"


def test_the_watermark_is_in_the_page_header_so_every_page_shows_it() -> None:
    document = Document(io.BytesIO(markdown_to_docx(SAMPLE, title="NDA")))

    headers = [p.text for s in document.sections for p in s.header.paragraphs]
    assert WATERMARK in headers


def test_the_disclaimer_is_in_the_body_not_only_the_metadata() -> None:
    document = Document(io.BytesIO(markdown_to_docx(SAMPLE, title="NDA")))
    text = "\n".join(p.text for p in document.paragraphs)

    assert DISCLAIMER in text
    assert "not legal advice" in text
    assert "must not be executed" in text


def test_headings_survive_as_headings_not_hashes() -> None:
    document = Document(io.BytesIO(markdown_to_docx(SAMPLE, title="NDA")))

    styles = {p.style.name for p in document.paragraphs if p.text == "Confidentiality"}
    assert styles and all(s.startswith("Heading") for s in styles)
    assert not any(p.text.startswith("#") for p in document.paragraphs)


def test_core_properties_are_fixed_so_the_document_does_not_vary() -> None:
    document = Document(io.BytesIO(markdown_to_docx(SAMPLE, title="NDA")))
    props = document.core_properties

    assert props.title == "NDA"
    assert props.created == props.modified
    assert props.created.year == 2000


def test_the_provenance_appendix_names_every_clause_and_its_version() -> None:
    from backend.clauselib.loader import clauses_for
    from backend.invariants.render import render_clause

    values = {
        "disclosing_party": "ABC Pvt Ltd",
        "receiving_party": "XYZ Pvt Ltd",
        "duration_years": "3",
        "effective_date": "1 August 2026",
        "term_end_date": "1 August 2029",
        "governing_law_country": "India",
        "jurisdiction_city": "Mumbai",
        "disclosing_signatory": "Jane Rao",
        "receiving_signatory": "Sam Patel",
    }
    clauses = [render_clause(c, {k: values[k] for k in c.variables}) for c in clauses_for("nda")]

    document = Document(io.BytesIO(markdown_to_docx(SAMPLE, title="NDA", clauses=clauses)))
    text = "\n".join(p.text for p in document.paragraphs)

    assert "Appendix: Clause Provenance" in text
    assert "nda.confidentiality@1" in text
    for clause in clauses:
        assert clause.source_sha[:12] in text


# ------------------------------------------------------------------------------ storage


def test_local_storage_refuses_a_key_that_could_escape_its_root(tmp_path: Path) -> None:
    storage = LocalStorage(tmp_path)
    for bad in ("../escape.docx", "sub/dir.docx", "..\\win.docx"):
        with pytest.raises(ValueError, match="illegal storage key"):
            storage.put(bad, b"x")


def test_local_storage_round_trips(tmp_path: Path) -> None:
    storage = LocalStorage(tmp_path)
    key = storage.put("a.docx", b"payload")

    assert storage.exists(key)
    assert storage.get(key) == b"payload"


# ------------------------------------------------------------------------------- the gate


@pytest_asyncio.fixture
async def ctx() -> AsyncIterator[RunContext]:
    engine = create_async_engine(settings.database_url)
    factory = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
    cid = uuid.uuid4()
    async with factory() as s:
        s.add(Contract(id=cid, contract_type="nda", request="Draft an NDA"))
        await s.commit()
    try:
        yield RunContext(contract_id=cid, session_factory=factory, contract_type="nda")
    finally:
        async with factory() as s:
            await s.execute(delete(Contract).where(Contract.id == cid))
            await s.commit()
        await engine.dispose()


async def _add_version(ctx: RunContext, attempt: int, markdown: str, score: int) -> uuid.UUID:
    async with ctx.session_factory() as s:
        version = ContractVersion(
            contract_id=ctx.contract_id,
            attempt=attempt,
            path=f"draft_v{attempt}.md",
            markdown=markdown,
            score=score,
        )
        s.add(version)
        await s.commit()
        return version.id


async def _export(ctx: RunContext, version_id: str) -> str:
    return str(
        await export_docx.on_invoke_tool(
            tool_ctx(ctx, "export_docx"), f'{{"contract_version_id": "{version_id}"}}'
        )
    )


async def _finalize(ctx: RunContext) -> None:
    await finalize_contract.on_invoke_tool(tool_ctx(ctx, "finalize_contract"), "{}")


async def test_a_finalized_version_exports(ctx: RunContext) -> None:
    await render_clauses_for(ctx)
    version_id = await _add_version(ctx, 1, await faithful_markdown(ctx), 95)
    await _finalize(ctx)

    result = await _export(ctx, str(version_id))

    assert "Exported as DOCX" in result
    assert "DRAFT — FOR LEGAL REVIEW" in result

    async with ctx.session_factory() as s:
        export = (
            await s.execute(select(Export).where(Export.contract_version_id == version_id))
        ).scalar_one()

    assert export.format == "docx"
    assert export.size_bytes > 0
    assert len(export.sha256) == 64
    assert LocalStorage(Path(settings.storage_dir)).exists(export.storage_key)


async def test_an_unfinalized_draft_cannot_be_exported(ctx: RunContext) -> None:
    """The version exists, the markdown is perfect, and it still cannot become a document."""
    await render_clauses_for(ctx)
    version_id = await _add_version(ctx, 1, await faithful_markdown(ctx), 95)

    result = await _export(ctx, str(version_id))

    assert "was never finalized" in result
    assert "Call finalize_contract first" in result

    async with ctx.session_factory() as s:
        rows = (
            (await s.execute(select(Export).where(Export.contract_version_id == version_id)))
            .scalars()
            .all()
        )
    assert rows == [], "no document may exist for an unfinalized version"


async def test_a_blocked_draft_can_never_be_exported(ctx: RunContext) -> None:
    """finalize refuses it, so no id exists to export. The two gates compose."""
    await render_clauses_for(ctx)
    version_id = await _add_version(ctx, 1, "# NDA\n\nTrust me.", 100)

    await _finalize(ctx)  # refuses
    result = await _export(ctx, str(version_id))

    assert "was never finalized" in result


@pytest.mark.parametrize(
    "bad", ["draft_v1.md", "final.md", "not-a-uuid", "", "'; DROP TABLE exports; --"]
)
async def test_export_refuses_anything_that_is_not_a_version_id(ctx: RunContext, bad: str) -> None:
    """There is no overload taking Markdown, and no path taking a workspace file."""
    result = await _export(ctx, bad)
    assert "is not a contract version id" in result


async def test_export_refuses_a_version_belonging_to_another_contract(ctx: RunContext) -> None:
    other = uuid.uuid4()
    async with ctx.session_factory() as s:
        s.add(Contract(id=other, contract_type="nda", request="other"))
        await s.commit()
        version = ContractVersion(
            contract_id=other, attempt=1, path="draft_v1.md", markdown=SAMPLE, score=99
        )
        s.add(version)
        await s.commit()
        stolen = version.id

    try:
        result = await _export(ctx, str(stolen))
        assert "no contract version" in result
    finally:
        async with ctx.session_factory() as s:
            await s.execute(delete(Contract).where(Contract.id == other))
            await s.commit()


async def test_export_refuses_an_unknown_version_id(ctx: RunContext) -> None:
    assert "no contract version" in await _export(ctx, str(uuid.uuid4()))


async def test_exporting_twice_reuses_the_same_document(ctx: RunContext) -> None:
    """The bytes are stable, so a second export is the same document, not a second one."""
    await render_clauses_for(ctx)
    version_id = await _add_version(ctx, 1, await faithful_markdown(ctx), 95)
    await _finalize(ctx)

    first = await _export(ctx, str(version_id))
    second = await _export(ctx, str(version_id))

    assert first == second
    async with ctx.session_factory() as s:
        rows = (
            (await s.execute(select(Export).where(Export.contract_version_id == version_id)))
            .scalars()
            .all()
        )
    assert len(rows) == 1


async def test_a_needs_review_export_says_so(ctx: RunContext) -> None:
    await render_clauses_for(ctx)
    version_id = await _add_version(ctx, 1, await faithful_markdown(ctx), 80)
    await _finalize(ctx)

    result = await _export(ctx, str(version_id))
    assert "needs_human_review" in result
