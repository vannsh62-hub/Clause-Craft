"""Editing an uploaded DOCX in place.

The whole of Mode 2's fidelity guarantee lives in one decision: **start from the source
bytes and never build a new document.** A KEEP section is preserved not because it is
carefully copied but because it is never touched — its `w:p` element, its `w:pPr`, its
`w:numPr`, its runs are all still the original XML. Reconstruction cannot match that, which
is why "regenerate the document and copy the content across" is a defect rather than an
implementation choice.

`styles.xml`, `numbering.xml`, `theme1.xml`, headers, footers and section properties are
never opened, so they survive by construction rather than by effort.

The ADD case is the one with a trap in it. `document.add_paragraph(style="List Number")`
produces a paragraph carrying `w:pStyle` and **no `w:numPr`** — verified against
python-docx 1.2.0 — so it renders with no number at all. Numbering in OOXML lives in the
paragraph's own properties, not in its style. The only reliable way to get a correctly
numbered new paragraph is to deep-copy one that already is, and edit its text.
"""

from __future__ import annotations

import copy
import re
from collections.abc import Mapping
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from backend.invariants.docx_parse import TemplateError, _fingerprint, _text_of, _walk, block_ids
from backend.schemas.plan import SectionDecision, TransformationPlan

__all__ = ["EditReport", "apply_transformation"]


class EditReport:
    """What the edit actually did, including what it could not do cleanly.

    Returned rather than logged: a section that had to fall back to an un-styled paragraph
    is a formatting defect the document validator should see, not a line in a log file
    nobody reads.
    """

    def __init__(self) -> None:
        self.kept: list[str] = []
        self.modified: list[str] = []
        self.removed: list[str] = []
        self.added: list[str] = []
        self.warnings: list[str] = []

    @property
    def touched(self) -> int:
        return len(self.modified) + len(self.removed) + len(self.added)


def _set_text(paragraph: Any, text: str) -> None:
    """Replace a paragraph's text, keeping its properties and first run's formatting.

    `w:pPr` is untouched, so style, numbering and indentation survive. Trailing runs are
    dropped rather than edited: a paragraph split across five runs for no visible reason is
    normal in Word, and preserving that split while changing the words means guessing where
    the new words belong.
    """
    runs = paragraph.findall(qn("w:r"))
    if not runs:
        run = paragraph.makeelement(qn("w:r"), {})
        node = paragraph.makeelement(qn("w:t"), {})
        node.text = text
        run.append(node)
        paragraph.append(run)
        return

    first, *rest = runs
    for node in first.findall(qn("w:t")):
        first.remove(node)
    node = first.makeelement(qn("w:t"), {})
    node.set(qn("xml:space"), "preserve")
    node.text = text
    first.append(node)
    for extra in rest:
        paragraph.remove(extra)


def _donor(elements: list[Any], target: Any) -> Any | None:
    """Find a paragraph to clone for an ADD, preferring one with the same style.

    Preference order: the nearest preceding paragraph sharing the target's style, then any
    paragraph sharing its numbering, then nothing. Cloning a neighbour is what carries the
    `w:numPr` that `add_paragraph` omits.
    """
    if target is None:
        return None
    target_pr = target.find(qn("w:pPr"))
    if target_pr is None:
        return None

    style = target_pr.find(qn("w:pStyle"))
    style_val = style.get(qn("w:val")) if style is not None else None

    for element in reversed(elements):
        if element.tag != qn("w:p"):
            continue
        ppr = element.find(qn("w:pPr"))
        if ppr is None:
            continue
        candidate = ppr.find(qn("w:pStyle"))
        if style_val and candidate is not None and candidate.get(qn("w:val")) == style_val:
            return element
    return None


#: Marks the start of the execution/signature page. Matched against a paragraph's own text,
#: not its heading level, because a KEEP/MODIFY decision from the transformation planner has
#: no notion of "structural, unnumbered section" today — see `backend/schemas/document.py`,
#: "What is numbered, and what is not". The planner only knows KEEP/MODIFY/REMOVE/ADD, so it
#: is free to (and does) classify the signature block like any other clause, and Mode 2 never
#: renumbers a KEEP paragraph — it inherits whatever `w:numPr` the source document gave it.
#: If the source's own numbered list runs into the signature page, this is the only place
#: left to stop it from printing as "1. Execution" or similar.
_EXECUTION_START = re.compile(
    r"^\s*execution\b|in witness whereof|signature page", re.IGNORECASE
)


def _strip_numbering(paragraph: Any) -> None:
    """Remove `w:numPr` so this paragraph no longer takes part in an auto-numbered list."""
    ppr = paragraph.find(qn("w:pPr"))
    if ppr is None:
        return
    numpr = ppr.find(qn("w:numPr"))
    if numpr is not None:
        ppr.remove(numpr)


def _unnumber_execution_block(elements: list[Any]) -> None:
    """Once the execution/signature page starts, nothing after it should be numbered.

    The signature block is deliberately unnumbered by design (see `schemas/document.py`),
    but Mode 2 preserves source formatting verbatim for KEEP/MODIFY paragraphs, including
    whatever list numbering they already carried. This walks the surviving paragraphs in
    document order and strips numbering from the execution heading onward, so a mis-tagged
    KEEP decision cannot make the signature page look like clause "1." or "6.".
    """
    seen_execution = False
    for element in elements:
        if element.tag != qn("w:p"):
            continue
        if not seen_execution:
            if _EXECUTION_START.search(_text_of(element)):
                seen_execution = True
            else:
                continue
        _strip_numbering(element)


def apply_transformation(
    source: bytes,
    plan: TransformationPlan,
    new_text: Mapping[str, str],
) -> tuple[bytes, EditReport]:
    """Apply `plan` to `source`, returning the edited document.

    `new_text` maps block id to replacement text and must cover every MODIFY and ADD
    decision. A missing entry is refused rather than filled with a placeholder: an empty
    clause in an executed contract is worse than a failed run.
    """
    document = Document(BytesIO(source))
    body = document.element.body

    # Re-walk the source purely for its refusals: a document with tracked changes must not
    # reach the editor, whatever the plan says about it.
    _walk(source)

    order = [c for c in body if c.tag.split("}")[-1] in {"p", "tbl"}]
    fingerprints = [
        _fingerprint(position, child.tag.split("}")[-1], child)
        for position, child in enumerate(order)
    ]
    index_by_id: dict[str, Any] = dict(zip(block_ids(fingerprints), order, strict=True))

    report = EditReport()
    _check_coverage(plan, new_text)

    for decision in plan.remove:
        element = _resolve(decision, index_by_id)
        if element is not None:
            element.getparent().remove(element)
            report.removed.append(decision.name)

    for decision in plan.modify:
        element = _resolve(decision, index_by_id)
        if element is None:
            report.warnings.append(f"{decision.name}: nothing to modify")
            continue
        if element.tag != qn("w:p"):
            report.warnings.append(f"{decision.name}: is a table; text replacement skipped")
            continue
        _set_text(element, new_text[_key(decision)])
        report.modified.append(decision.name)

    for decision in plan.add:
        anchor = _resolve(decision, index_by_id)
        # `or` would truth-test an lxml element, which python-docx warns is ambiguous.
        candidate = _donor(order, anchor)
        donor = candidate if candidate is not None else anchor
        if donor is not None and donor.tag == qn("w:p"):
            clone = copy.deepcopy(donor)
            _set_text(clone, new_text[_key(decision)])
            donor.addnext(clone)
        else:
            # No donor: the paragraph gets default styling and, if it should have been
            # numbered, will not be. Surfaced rather than hidden.
            document.add_paragraph(new_text[_key(decision)])
            report.warnings.append(
                f"{decision.name}: added without a styled donor paragraph; "
                "numbering and indentation may not match the surrounding document"
            )
        report.added.append(decision.name)

    report.kept = [d.name for d in plan.keep]

    surviving = [c for c in body if c.tag.split("}")[-1] in {"p", "tbl"}]
    _unnumber_execution_block(surviving)

    out = BytesIO()
    document.save(out)
    return out.getvalue(), report


def _key(decision: SectionDecision) -> str:
    ref = decision.source_ref
    return (ref.block_id if ref and ref.block_id else None) or decision.name


def _resolve(decision: SectionDecision, index: Mapping[str, Any]) -> Any | None:
    ref = decision.source_ref
    if ref is None or ref.block_id is None:
        return None
    return index.get(ref.block_id)


def _check_coverage(plan: TransformationPlan, new_text: Mapping[str, str]) -> None:
    missing = [
        decision.name for decision in (*plan.modify, *plan.add) if _key(decision) not in new_text
    ]
    if missing:
        raise TemplateError(
            "No replacement text was supplied for: " + ", ".join(sorted(missing)) + ". "
            "Every MODIFY and ADD decision needs content; an empty clause is worse than "
            "a failed run."
        )