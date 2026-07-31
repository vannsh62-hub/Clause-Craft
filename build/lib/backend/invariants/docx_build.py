"""Building a contract document. Deterministic, model-free, and shaped like a contract.

`markdown_to_docx` renders headings and paragraphs, which is why our generated drafts came
out flat: no title page, no party block, no clause numbers, no execution page. This renders a
`ContractDocument` instead, and applies the conventions a legal document actually follows.

## Numbering is Word's job, not the model's

The operative clauses carry `w:numPr` and reference a multilevel numbering definition
injected into the package's `numbering.xml`: `1.` → `1.1` → `(a)` → `(i)`. Word computes the
labels. Nothing types "1.1" into text, so inserting a clause renumbers everything below it
automatically — the failure that makes hand-numbered drafts fall apart on the first edit.

The preamble, recitals and signature block are deliberately **not** numbered. A numbered
"WHEREAS" is the clearest tell that a document was assembled by a machine.

## The text form is derived from the same object

`document_to_text` renders the same tree with explicit numbers. The validation gates and the
UI preview read that text, so deriving both from one `ContractDocument` is what stops the
document being validated from drifting away from the document being shipped.

Imports neither `agents` nor `openai`, and must not — `test_invariants_are_llm_free.py`
globs this directory.
"""

from __future__ import annotations

import io

from docx import Document as new_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from backend.invariants.export import EPOCH, normalise_zip
from backend.schemas.document import ClauseNode, ContractDocument

__all__ = ["document_to_text", "render_contract"]

#: Ids for the numbering definition we inject. High enough not to collide with the nine
#: definitions python-docx's default template already ships.
_ABSTRACT_NUM_ID = 9100
_NUM_ID = 9100

#: One level per depth: 1. / 1.1 / (a) / (i). Deeper than four is not contract drafting.
_LEVELS: tuple[tuple[str, str], ...] = (
    ("decimal", "%1."),
    ("decimal", "%1.%2"),
    ("lowerLetter", "(%3)"),
    ("lowerRoman", "(%4)"),
)

_INDENT_TWIPS = 425  # 0.75 cm per level


def _install_numbering(document: Document) -> None:
    """Add the legal multilevel numbering definition to the package.

    `w:abstractNum` describes the label format per level; `w:num` is the concrete instance
    paragraphs point at. Both are required — a `w:numPr` referencing an id with no `w:num`
    renders as no numbering at all rather than as an error, which is the kind of failure that
    reaches a reviewer before it reaches a test.
    """
    numbering = document.part.numbering_part.element

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(_ABSTRACT_NUM_ID))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for depth, (fmt, text) in enumerate(_LEVELS):
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), str(depth))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)

        justify = OxmlElement("w:lvlJc")
        justify.set(qn("w:val"), "left")
        level.append(justify)

        # Indent each level, and hang the text off the label so wrapped lines line up.
        p_pr = OxmlElement("w:pPr")
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(_INDENT_TWIPS * (depth + 1)))
        indent.set(qn("w:hanging"), str(_INDENT_TWIPS))
        p_pr.append(indent)
        level.append(p_pr)

        abstract.append(level)

    # `w:abstractNum` must precede `w:num` in the part, so insert at the front and append the
    # instance at the end.
    numbering.insert(0, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(_NUM_ID))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(_ABSTRACT_NUM_ID))
    num.append(ref)
    numbering.append(num)


def _number(paragraph: object, depth: int) -> None:
    """Attach the numbering definition to a paragraph at `depth`."""
    p_pr = paragraph._p.get_or_add_pPr()  # type: ignore[attr-defined]
    num_pr = OxmlElement("w:numPr")

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(depth, len(_LEVELS) - 1)))
    num_pr.append(ilvl)

    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(_NUM_ID))
    num_pr.append(num_id)

    p_pr.append(num_pr)


def _add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(15)


def _add_clause(document: Document, clause: ClauseNode, depth: int) -> None:
    """One clause and its children. Heading and body share the number.

    The heading paragraph carries the numbering; the body is a plain indented paragraph
    beneath it, so the clause reads as "1. Confidentiality" followed by its text rather than
    as two separately numbered items.
    """
    heading = document.add_paragraph()
    _number(heading, depth)
    run = heading.add_run(clause.heading)
    run.bold = True

    if clause.text.strip():
        for block in clause.text.strip().split("\n\n"):
            body = document.add_paragraph(block.strip())
            body.paragraph_format.left_indent = Pt(18 * (depth + 1))

    for child in clause.children:
        _add_clause(document, child, depth + 1)


def _add_signature_block(document: Document, doc: ContractDocument) -> None:
    """The execution page. Unnumbered, and on a page of its own.

    A signature block that shares a page with the last clause is the other common tell of a
    generated document; contracts execute on a fresh page.
    """
    if not doc.signatories:
        return

    document.add_page_break()  # type: ignore[no-untyped-call]

    if doc.execution_note:
        document.add_paragraph(doc.execution_note)

    for signatory in doc.signatories:
        document.add_paragraph()
        label = document.add_paragraph()
        run = label.add_run(f"For and on behalf of {signatory.name}")
        run.bold = True
        if signatory.role:
            document.add_paragraph(f"({signatory.role})")
        for line in ("By: ______________________________", "Name:", "Title:", "Date:"):
            document.add_paragraph(line)


def render_contract(doc: ContractDocument) -> bytes:
    """Render a `ContractDocument` to DOCX. The same document always yields the same bytes."""
    document = new_document()

    document.core_properties.title = doc.title
    document.core_properties.author = "AI Contract Drafting Platform"
    document.core_properties.created = EPOCH
    document.core_properties.modified = EPOCH
    document.core_properties.revision = 1

    _install_numbering(document)

    _add_title(document, doc.title)

    if doc.preamble:
        document.add_paragraph()
        document.add_paragraph(doc.preamble)

    if doc.recitals:
        document.add_paragraph()
        for recital in doc.recitals:
            document.add_paragraph(recital)

    if doc.clauses:
        document.add_paragraph()
        for clause in doc.clauses:
            _add_clause(document, clause, 0)

    _add_signature_block(document, doc)

    buffer = io.BytesIO()
    document.save(buffer)
    return normalise_zip(buffer.getvalue())


def _label(counters: list[int], depth: int) -> str:
    """The label Word will render, computed here for the text form.

    Kept in step with `_LEVELS` by construction: decimal, dotted decimal, letter, roman.
    """
    if depth == 0:
        return f"{counters[0]}."
    if depth == 1:
        return f"{counters[0]}.{counters[1]}"
    if depth == 2:
        return f"({chr(ord('a') + counters[2] - 1)})"
    roman = ("i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x")
    index = min(counters[3], len(roman)) - 1
    return f"({roman[index]})"


def _clause_text(clause: ClauseNode, depth: int, counters: list[int], out: list[str]) -> None:
    while len(counters) <= depth:
        counters.append(0)
    counters[depth] += 1
    del counters[depth + 1 :]

    out.append(f"{_label(counters, depth)} {clause.heading}".strip())
    if clause.text.strip():
        out.append(clause.text.strip())
    out.append("")

    for child in clause.children:
        _clause_text(child, depth + 1, counters, out)


def document_to_text(doc: ContractDocument) -> str:
    """The same document as text, with numbers written out.

    This is what the validation gates read and what the UI previews, so it must contain
    everything the DOCX contains — a gate that cannot see the signature block cannot tell you
    it is missing.
    """
    out: list[str] = [f"# {doc.title}", ""]

    if doc.preamble:
        out.extend([doc.preamble, ""])

    if doc.recitals:
        out.extend([*doc.recitals, ""])

    counters: list[int] = []
    for clause in doc.clauses:
        _clause_text(clause, 0, counters, out)

    if doc.signatories:
        if doc.execution_note:
            out.extend([doc.execution_note, ""])
        for signatory in doc.signatories:
            out.append(f"For and on behalf of {signatory.name}")
            if signatory.role:
                out.append(f"({signatory.role})")
            out.extend(["By: ______________________________", "Name:", "Title:", "Date:", ""])

    return "\n".join(out).strip() + "\n"
