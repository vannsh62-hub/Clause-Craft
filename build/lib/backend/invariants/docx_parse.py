"""Parsing a DOCX into structure.

Deterministic, no model. Parsing a document has a correct answer, and the things that go
wrong here — a lost table row, a dropped numbering level — are the things a lawyer notices
first.

Three decisions, each of which was arrived at by looking at what `python-docx` actually
does rather than what it appears to do:

**Walk the body, not `document.paragraphs`.** `.paragraphs` silently omits paragraphs
inside tables. That is exactly where numbered schedules, fee tables and annexures live, so
a parser built on it under-reports precisely the content most likely to matter. Verified:
a paragraph placed in a table cell does not appear in `.paragraphs`.

**Fingerprint the formatting; do not hash the XML.** `python-docx` reserialises the tree on
save — attribute order, namespace declarations and whitespace shift even for a load then
save with no edits — so comparing package bytes is flaky. Spec 05 §2 asks for
"byte-comparable" output, which is not achievable; what is achievable, and is what
paragraph-style fidelity actually means, is that style ids, numbering identity and run
properties are unchanged.

**Refuse tracked changes.** Accepting them silently mutates a legal document with no audit
trail. Rejecting them silently discards negotiated edits. Either way the fidelity guarantee
becomes unverifiable, because the document that was reviewed is not the document that was
parsed. Refusing is the only honest option, and the user can resolve them in Word in
seconds.
"""

from __future__ import annotations

import hashlib
import re
from collections import Counter
from collections.abc import Sequence
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from backend.schemas.errors import ContractToolError
from backend.schemas.template import (
    BlockFingerprint,
    FormattingManifest,
    NumberingScheme,
    ParsedTable,
    Placeholder,
    TemplateObject,
)

__all__ = [
    "TemplateError",
    "block_ids",
    "block_texts",
    "parse_docx",
    "style_fingerprint",
]

#: `{{ name }}`, `[NAME]`, `<INSERT NAME>` — the three conventions seen in real templates.
_PLACEHOLDER = re.compile(r"\{\{[^}]+\}\}|\[[A-Z][A-Z0-9 _/-]{2,}\]|<[^>]{3,}>")

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class TemplateError(ContractToolError):
    """A document this parser will not model.

    Raised rather than degraded. A refusal is a support ticket; a silently mangled
    executed contract is a lawsuit.
    """


def _text_of(element: Any) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _run_props(paragraph: Any) -> tuple[str, ...]:
    """A compact digest of each run's character formatting.

    Enough to notice that bold went missing or the font changed; not so detailed that an
    irrelevant reserialisation counts as a difference.
    """
    props: list[str] = []
    for run in paragraph.findall(qn("w:r")):
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            props.append("-")
            continue
        flags = [tag for tag in ("w:b", "w:i", "w:u", "w:strike") if rpr.find(qn(tag)) is not None]
        size = rpr.find(qn("w:sz"))
        if size is not None:
            flags.append(f"sz{size.get(qn('w:val'))}")
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is not None and fonts.get(qn("w:ascii")):
            flags.append(str(fonts.get(qn("w:ascii"))))
        props.append("|".join(flags) or "-")
    return tuple(props)


def _paragraph_fingerprint(index: int, paragraph: Any) -> BlockFingerprint:
    ppr = paragraph.find(qn("w:pPr"))
    style_id = num_id = ilvl = outline = None

    if ppr is not None:
        style = ppr.find(qn("w:pStyle"))
        if style is not None:
            style_id = style.get(qn("w:val"))

        numpr = ppr.find(qn("w:numPr"))
        if numpr is not None:
            num = numpr.find(qn("w:numId"))
            lvl = numpr.find(qn("w:ilvl"))
            num_id = int(num.get(qn("w:val"))) if num is not None else None
            ilvl = int(lvl.get(qn("w:val"))) if lvl is not None else None

        outline_el = ppr.find(qn("w:outlineLvl"))
        if outline_el is not None:
            outline = int(outline_el.get(qn("w:val")))

    text = _text_of(paragraph)
    return BlockFingerprint(
        index=index,
        kind="paragraph",
        style_id=style_id,
        num_id=num_id,
        ilvl=ilvl,
        outline_level=outline,
        run_props=_run_props(paragraph),
        text_sha=hashlib.sha256(" ".join(text.split()).encode("utf-8")).hexdigest()[:16],
    )


def _table_shape(table: Any) -> tuple[int, int, tuple[tuple[int, int], ...]]:
    rows = table.findall(qn("w:tr"))
    widths = [len(row.findall(qn("w:tc"))) for row in rows]
    merged: list[tuple[int, int]] = []
    for r, row in enumerate(rows):
        for c, cell in enumerate(row.findall(qn("w:tc"))):
            tc_pr = cell.find(qn("w:tcPr"))
            if tc_pr is None:
                continue
            if tc_pr.find(qn("w:gridSpan")) is not None or tc_pr.find(qn("w:vMerge")) is not None:
                merged.append((r, c))

    return len(rows), (max(widths) if widths else 0), tuple(merged)


def _table_fingerprint(index: int, table: Any) -> BlockFingerprint:
    text = _text_of(table)
    return BlockFingerprint(
        index=index,
        kind="table",
        text_sha=hashlib.sha256(" ".join(text.split()).encode("utf-8")).hexdigest()[:16],
    )


def block_ids(fingerprints: Sequence[BlockFingerprint]) -> list[str]:
    """Stable identities for a document's blocks.

    **Position is deliberately not part of the identity.** An id containing the block index
    is stable across a load/save cycle but not across an *edit*: removing one paragraph
    renumbers every block after it, so every KEEP decision downstream of a REMOVE would
    point at nothing and the fidelity check would report the untouched blocks as deleted.
    Since surviving edits is the entire purpose of these ids, identity is content plus
    style.

    That alone is not unique — contracts genuinely repeat paragraphs, most often empty ones
    and "Intentionally left blank" — so identical blocks are distinguished by how many
    identical blocks preceded them. That ordinal is stable as long as the duplicates are
    not reordered among themselves, which no transformation does.
    """
    seen: Counter[str] = Counter()
    ids: list[str] = []
    for fingerprint in fingerprints:
        seed = f"{fingerprint.style_id or ''}|{fingerprint.text_sha}"
        occurrence = seen[seed]
        seen[seed] += 1
        ids.append(hashlib.sha256(f"{seed}|{occurrence}".encode()).hexdigest()[:16])
    return ids


def _reject_tracked_changes(body: Any) -> None:
    if body.find(f".//{{{_W}}}ins") is not None or body.find(f".//{{{_W}}}del") is not None:
        raise TemplateError(
            "This document has tracked changes. Accept or reject them in Word, then "
            "re-upload. Editing a document whose revisions are unresolved would either "
            "silently accept someone's edits or silently discard them."
        )


def _walk(data: bytes) -> tuple[Any, list[tuple[int, str, Any]]]:
    """Return the document and its body blocks, in document order."""
    document = Document(BytesIO(data))
    body = document.element.body
    _reject_tracked_changes(body)

    blocks: list[tuple[int, str, Any]] = []
    index = 0
    for child in body:
        tag = child.tag.split("}")[-1]
        if tag in {"p", "tbl"}:
            blocks.append((index, tag, child))
            index += 1
    return document, blocks


def _fingerprint(index: int, tag: str, element: Any) -> BlockFingerprint:
    if tag == "p":
        return _paragraph_fingerprint(index, element)
    return _table_fingerprint(index, element)


def style_fingerprint(data: bytes) -> tuple[BlockFingerprint, ...]:
    """The formatting shape of a document, for comparison before and after an edit."""
    _, blocks = _walk(data)
    return tuple(_fingerprint(index, tag, element) for index, tag, element in blocks)


def block_texts(data: bytes) -> dict[str, str]:
    """Block id → text, for the agents that need to read the document's content.

    One walk. Pairing two independently-produced sequences by position would be correct
    right up until one of them started filtering.
    """
    _, blocks = _walk(data)
    fingerprints = [_fingerprint(index, tag, element) for index, tag, element in blocks]
    return {
        identity: _text_of(element)
        for identity, (_, _, element) in zip(block_ids(fingerprints), blocks, strict=True)
    }


def parse_docx(data: bytes, *, filename: str, storage_key: str) -> TemplateObject:
    """Parse an uploaded document into a `TemplateObject`.

    The bytes themselves are not carried: `workspace_files.content` is a text column, so
    the binary lives in object storage and is referenced by `storage_key`.
    """
    document, blocks = _walk(data)

    fingerprints = [_fingerprint(index, tag, element) for index, tag, element in blocks]
    identities = block_ids(fingerprints)

    tables: list[ParsedTable] = []
    placeholders: list[Placeholder] = []

    for identity, (_, tag, element) in zip(identities, blocks, strict=True):
        if tag == "tbl":
            rows, columns, merged = _table_shape(element)
            tables.append(ParsedTable(block_id=identity, rows=rows, columns=columns, merged=merged))
        for token in _PLACEHOLDER.findall(_text_of(element)):
            placeholders.append(Placeholder(token=token, block_id=identity, name=_name_of(token)))

    num_ids = tuple(sorted({f.num_id for f in fingerprints if f.num_id is not None}))
    depths = [f.ilvl for f in fingerprints if f.ilvl is not None]

    sections = document.element.body.findall(f".//{{{_W}}}sectPr")
    return TemplateObject(
        storage_key=storage_key,
        sha256=hashlib.sha256(data).hexdigest(),
        size_bytes=len(data),
        filename=filename,
        formatting=FormattingManifest(
            blocks=tuple(fingerprints),
            numbering=NumberingScheme(
                num_ids=num_ids,
                max_depth=max(depths) + 1 if depths else 0,
            ),
            tables=tuple(tables),
            has_headers=bool(document.sections and document.sections[0].header.paragraphs),
            has_footers=bool(document.sections and document.sections[0].footer.paragraphs),
            # One sectPr always closes the body; extras are real breaks.
            section_breaks=max(len(sections) - 1, 0),
        ),
        placeholders=tuple(placeholders),
    )


def _name_of(token: str) -> str | None:
    stripped = token.strip("{}[]<> ").strip()
    return stripped or None
