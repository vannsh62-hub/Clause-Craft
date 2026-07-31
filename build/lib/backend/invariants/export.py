"""Markdown to DOCX. Deterministic, and model-free.

Built programmatically rather than by filling a `.docx` template, because `python-docx`
cannot fill templates — it constructs documents. There is no `templates/contract.docx`.

**Byte stability took work.** `python-docx` stamps every entry in the OPC zip with the wall
clock, so two saves of identical content differ. The content is identical; the container is
not. `normalise_zip` rewrites the archive with fixed timestamps, sorted entries and a fixed
compression level, so the same Markdown always produces the same bytes. That is what makes an
export reproducible, and what lets `exports.sha256` mean something.

Imports neither `agents` nor `openai`.
"""

from __future__ import annotations

import hashlib
import io
import zipfile
from collections.abc import Sequence
from datetime import datetime

from docx import Document as new_document
from docx.document import Document  # the class; `docx.Document` is a factory function
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from backend.schemas.clause import RenderedClause

__all__ = [
    "DISCLAIMER",
    "EPOCH",
    "WATERMARK",
    "docx_sha256",
    "markdown_to_docx",
    "normalise_zip",
]

WATERMARK = "DRAFT — FOR LEGAL REVIEW"

DISCLAIMER = (
    "This document was assembled by an automated system from a library of counsel-approved "
    "clauses. It is a draft. It is not legal advice, it has not been reviewed by a lawyer, "
    "and it must not be executed until it has been."
)

#: A fixed instant, so the document properties do not vary between two exports of one draft.
EPOCH = datetime(2000, 1, 1)  # noqa: DTZ001 - deliberately naive and constant

#: DOS epoch, the earliest a zip timestamp can express.
_ZIP_EPOCH = (1980, 1, 1, 0, 0, 0)


def normalise_zip(raw: bytes) -> bytes:
    """Rewrite an OPC package so identical content yields identical bytes.

    `[Content_Types].xml` is written first, as the OPC specification expects; the rest follow
    in sorted order so the archive does not depend on dictionary iteration.
    """
    with zipfile.ZipFile(io.BytesIO(raw)) as source:
        entries = {name: source.read(name) for name in source.namelist()}

    ordered = sorted(entries, key=lambda n: (n != "[Content_Types].xml", n))

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as target:
        for name in ordered:
            info = zipfile.ZipInfo(name, date_time=_ZIP_EPOCH)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, entries[name])
    return out.getvalue()


def _add_watermark_header(document: Document) -> None:
    """A banner in the page header, repeated on every page.

    Not a diagonal WordArt watermark: that is VML in the header and renders inconsistently
    outside Word. A header banner is visible in Word and LibreOffice alike, which is what
    "the reviewer must see this" actually requires.
    """
    for section in document.sections:
        paragraph = section.header.paragraphs[0]
        paragraph.text = WATERMARK
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)


def _add_disclaimer(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(DISCLAIMER)
    run.italic = True
    run.font.size = Pt(9)


def _add_body(document: Document, markdown: str) -> None:
    """Headings by prefix; every other non-blank line is its own paragraph.

    Line-per-paragraph preserves the signature block, whose lines are single-newline
    separated and must not be reflowed into one another.
    """
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        else:
            document.add_paragraph(line.strip())


def _add_provenance(document: Document, clauses: Sequence[RenderedClause]) -> None:
    """Every clause, named and versioned. A reviewer can check any of them against the library."""
    if not clauses:
        return

    # `add_page_break` is unannotated in python-docx; everything else on Document is typed.
    document.add_page_break()  # type: ignore[no-untyped-call]
    document.add_heading("Appendix: Clause Provenance", level=2)
    document.add_paragraph(
        "Each clause below was reproduced verbatim from the approved clause library."
    )
    for clause in sorted(clauses, key=lambda c: c.order):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{clause.title} — ").bold = True
        run = paragraph.add_run(f"{clause.provenance} (sha256 {clause.source_sha[:12]})")
        run.font.size = Pt(9)


def markdown_to_docx(
    markdown: str,
    *,
    title: str,
    clauses: Sequence[RenderedClause] = (),
) -> bytes:
    """Render a finalized contract. The same Markdown always yields the same bytes."""
    document = new_document()

    document.core_properties.title = title
    document.core_properties.author = "AI Contract Drafting Platform"
    document.core_properties.comments = WATERMARK
    document.core_properties.created = EPOCH
    document.core_properties.modified = EPOCH
    document.core_properties.revision = 1

    _add_watermark_header(document)
    _add_disclaimer(document)
    _add_body(document, markdown)
    _add_provenance(document, clauses)

    buffer = io.BytesIO()
    document.save(buffer)
    return normalise_zip(buffer.getvalue())


def docx_sha256(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()
